import streamlit as st
import anthropic
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
from datetime import datetime

# Page config
st.set_page_config(
    page_title="AI Resume Tailor",
    page_icon="📄",
    layout="centered",
    initial_sidebar_state="collapsed"
)

# Path to your master resume
MASTER_RESUME_PATH = "master_resume.docx"

# Custom CSS - Force single card design
st.markdown("""
<style>
    /* Hide all Streamlit UI elements */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    .stDeployButton {display: none;}
    [data-testid="collapsedControl"] {display: none;}
    
    /* Purple gradient background - force on everything */
    .stApp,
    [data-testid="stAppViewContainer"],
    [data-testid="stHeader"],
    .main {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
    }
    
    /* Remove all default padding and margins */
    .main .block-container {
        padding: 2rem 1rem !important;
        max-width: 650px !important;
    }
    
    /* Single white card container - force all content into one */
    [data-testid="stVerticalBlock"] {
        background: white !important;
        border-radius: 24px !important;
        padding: 3rem 2.5rem !important;
        box-shadow: 0 25px 50px rgba(0, 0, 0, 0.25) !important;
        gap: 1.5rem !important;
    }
    
    /* Remove spacing between elements */
    [data-testid="stVerticalBlock"] > div {
        gap: 0.5rem !important;
    }
    
    /* Header styling */
    h1 {
        font-size: 2.5rem !important;
        font-weight: 700 !important;
        color: #1a202c !important;
        text-align: center !important;
        margin: 0 0 0.5rem 0 !important;
    }
    
    .subtitle {
        font-size: 1.1rem !important;
        color: #718096 !important;
        text-align: center !important;
        margin-bottom: 2rem !important;
    }
    
    /* Input labels */
    [data-testid="stVerticalBlock"] label {
        color: #2d3748 !important;
        font-weight: 600 !important;
        font-size: 0.9rem !important;
        margin-bottom: 0.5rem !important;
    }
    
    /* Text inputs - white with border */
    .stTextInput input {
        background-color: white !important;
        color: #1a202c !important;
        border: 2px solid #e2e8f0 !important;
        border-radius: 12px !important;
        padding: 0.75rem 1rem !important;
        font-size: 1rem !important;
    }
    
    .stTextInput input:focus {
        border-color: #667eea !important;
        box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1) !important;
    }
    
    /* Text area */
    .stTextArea textarea {
        background-color: white !important;
        color: #1a202c !important;
        border: 2px solid #e2e8f0 !important;
        border-radius: 12px !important;
        padding: 0.75rem 1rem !important;
        font-size: 1rem !important;
        min-height: 100px !important;
    }
    
    .stTextArea textarea:focus {
        border-color: #667eea !important;
        box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1) !important;
    }
    
    /* File uploader */
    [data-testid="stFileUploader"] {
        background-color: transparent !important;
    }
    
    [data-testid="stFileUploader"] section {
        background-color: #f7fafc !important;
        border: 2px dashed #cbd5e0 !important;
        border-radius: 12px !important;
        padding: 1.5rem !important;
    }
    
    [data-testid="stFileUploader"] button {
        background: #667eea !important;
        color: white !important;
        border: none !important;
        border-radius: 8px !important;
        padding: 0.5rem 1.5rem !important;
        font-weight: 600 !important;
    }
    
    [data-testid="stFileUploader"] small {
        color: #718096 !important;
    }
    
    /* Main button */
    .stButton button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
        color: white !important;
        border: none !important;
        border-radius: 12px !important;
        padding: 1rem 2rem !important;
        font-size: 1.1rem !important;
        font-weight: 600 !important;
        width: 100% !important;
        cursor: pointer !important;
        transition: transform 0.2s !important;
    }
    
    .stButton button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 10px 25px rgba(102, 126, 234, 0.3) !important;
    }
    
    /* Feature badges */
    .feature-grid {
        display: grid;
        grid-template-columns: repeat(2, 1fr);
        gap: 1rem;
        margin-top: 2rem;
        padding-top: 2rem;
        border-top: 2px solid #e2e8f0;
    }
    
    .feature-item {
        display: flex;
        align-items: center;
        gap: 0.75rem;
    }
    
    .feature-icon {
        width: 28px;
        height: 28px;
        background: #48bb78;
        color: white;
        border-radius: 50%;
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 16px;
        font-weight: bold;
        flex-shrink: 0;
    }
    
    .feature-text {
        color: #4a5568;
        font-size: 0.95rem;
        font-weight: 500;
    }
    
    /* Success/Error messages */
    .stSuccess, .stError, .stInfo {
        background-color: white !important;
        border-radius: 8px !important;
    }
    
    /* FORCE remove all animations and fading - nuclear option */
    *, *::before, *::after, div, p, span, h1, h2, h3, h4, h5, h6 {
        animation: none !important;
        animation-duration: 0s !important;
        animation-delay: 0s !important;
        transition: none !important;
        transition-duration: 0s !important;
        transition-delay: 0s !important;
        opacity: 1 !important;
        visibility: visible !important;
    }
    
    /* ONLY allow button hover transition */
    .stButton button {
        transition: transform 0.2s, box-shadow 0.2s !important;
    }
    
    /* Force all text elements to be immediately visible */
    [data-testid="stMarkdownContainer"],
    [data-testid="stMarkdownContainer"] p,
    [data-testid="stMarkdownContainer"] strong,
    .element-container,
    .stMarkdown {
        opacity: 1 !important;
        visibility: visible !important;
        animation: none !important;
        transition: none !important;
    }
    
    /* Progress bar styling */
    .stProgress > div > div {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
    }
    
    /* Force status messages to appear instantly */
    .stAlert, .stSuccess, .stError, .stInfo, .stWarning {
        opacity: 1 !important;
        visibility: visible !important;
        animation: none !important;
        transition: none !important;
    }
</style>
""", unsafe_allow_html=True)

def scrape_job_posting(url):
    """Scrape job posting content from URL"""
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()
        text = soup.get_text(separator='\n', strip=True)
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        return '\n'.join(lines)[:15000]
    except Exception as e:
        st.error(f"Error scraping: {str(e)}")
        return None

def analyze_job_posting(job_content, api_key):
    """Analyze job posting with Claude"""
    client = anthropic.Anthropic(api_key=api_key)
    prompt = f"""Analyze this job posting and extract requirements in a structured format.

JOB POSTING:
{job_content}

Extract and organize as follows:

1. JOB TITLE: [exact title from posting]

2. COMPANY & LOCATION: [company name and location]

3. MUST-HAVE REQUIREMENTS (Deal-breakers - usually 5-8 items):
   List the absolute requirements for the role:
   - Prior experience requirements
   - Technical skills that are mandatory
   - Industry experience required
   - Regulatory/compliance requirements
   - Education/certification requirements

4. KEY RESPONSIBILITIES (Top 7-10 main duties):
   List the primary job functions in order of importance

5. TECHNICAL SKILLS REQUIRED (8-15 items):
   - Software/systems
   - Laboratory equipment/instruments
   - Analytical techniques
   - Technical competencies

6. SOFT SKILLS & ATTRIBUTES:
   - Work style (autonomous, collaborative, etc.)
   - Personal qualities mentioned
   - Communication requirements

7. CRITICAL ATS KEYWORDS (20-30 keywords):
   List exact phrases and terms from the job posting that should appear in resume:
   - Job title variations
   - Technical terminology
   - Industry-specific terms
   - Required qualifications
   - Regulatory terms (GMP, GLP, FDA, etc.)
   - Action words used in posting

8. LOCATION DETAILS:
   - Work location
   - On-site/hybrid/remote
   - Any relocation notes

9. COMPANY BACKGROUND:
   - Industry
   - Company focus/mission if mentioned
   - Team/department details

Format your response clearly with headers and bullet points for easy parsing."""

    try:
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=3000,
            messages=[{"role": "user", "content": prompt}]
        )
        return message.content[0].text
    except Exception as e:
        st.error(f"Error analyzing job: {str(e)}")
        return None

def extract_resume_content(uploaded_file):
    """Extract resume content"""
    try:
        doc = Document(uploaded_file)
        return '\n'.join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
    except Exception as e:
        st.error(f"Error reading resume: {str(e)}")
        return None

def generate_tailored_resume(original_resume, job_analysis, additional_context, api_key):
    """Generate tailored resume using Claude"""
    client = anthropic.Anthropic(api_key=api_key)
    
    prompt = f"""You are an expert ATS resume writer and career strategist. Your job is to create a highly targeted resume that will get past ATS systems AND impress human hiring managers.

ORIGINAL RESUME:
{original_resume}

JOB REQUIREMENTS ANALYSIS:
{job_analysis}

ADDITIONAL CONTEXT:
{additional_context if additional_context else "None"}

====================
CRITICAL INSTRUCTIONS
====================

1. EXPERIENCE REORDERING (MOST IMPORTANT):
   - Analyze which past role is MOST RELEVANT to this specific job
   - Move the most relevant role to POSITION #1 in work experience, even if it's not the most recent
   - This is critical for ATS ranking and human reviewer attention
   - Keep dates accurate but reorder by strategic fit

2. PROFESSIONAL SUMMARY REWRITE:
   - Start with the EXACT job title from the posting (or closest match)
   - Include the top 5-7 required qualifications from the job description
   - Mention specific years of experience in the relevant specialty
   - If candidate has experience in the same city/region as job, mention it
   - Use confident, active language
   - NO generic phrases like "detail-oriented professional" or "experienced technician"
   - Make it laser-focused on THIS specific role

3. KEYWORD DISCIPLINE (CRITICAL):
   - ONLY use technical terms and keywords that appear in:
     a) The job description, OR
     b) The original resume
   - DO NOT add new technical terms, industry buzzwords, or specialty areas not in either document
   - DO NOT hallucinate skills like "microbiological analysis" or "food safety" unless in job posting
   - When in doubt, use the EXACT wording from the job description

4. ADDRESS EVERY JOB REQUIREMENT:
   For each requirement in the job posting:
   - If it's explicitly in the resume, emphasize it prominently
   - If it's IMPLIED by existing experience, rephrase bullets to make it explicit
   - If it's MISSING, create a new bullet that shows how existing experience relates to it
   
   Examples:
   - Job needs "batch release testing" but resume says "finished product testing" 
     → Change to "Execute final product testing to support batch release decisions"
   - Job needs "laboratory setup" but resume doesn't mention it explicitly
     → Add "Supported laboratory operationalization and equipment setup during facility expansion"

5. BULLET POINT PRIORITIZATION:
   Within each role, reorder bullets to put:
   - Most relevant to THIS job's requirements = FIRST (top 3 bullets)
   - Quantified achievements = MIDDLE
   - General responsibilities = LAST
   
   Each bullet should start with a strong action verb and include metrics when possible.

6. FORMATTING REQUIREMENTS:
   - NO markdown formatting (no #, ##, **, etc.)
   - Use UPPERCASE for section headers: PROFESSIONAL SUMMARY, WORK EXPERIENCE, EDUCATION
   - Use clear line breaks between sections
   - Bold job titles and company names (use proper Word formatting, not markdown)
   - Bullet points with actual bullet characters (•), not dashes or asterisks
   - Keep contact information in a single line at the top
   - Professional, clean formatting suitable for Word document conversion

7. LOCATION EMPHASIS:
   - If job location matches any past role location, emphasize this connection
   - Mention it in the summary if significant (e.g., "including Sydney-based pharmaceutical operations")

8. REMOVE FLUFF:
   - No "References available upon request" (assumed)
   - No generic skills that aren't job-specific
   - No outdated or irrelevant experience unless it shows progression
   - Focus on last 10-15 years of experience

9. ATS OPTIMIZATION:
   - Use standard section headers ATS systems recognize
   - Include job title keywords in summary
   - Mirror language from job description throughout
   - Avoid tables, text boxes, or complex formatting
   - Use standard job title formats

10. QUANTIFY EVERYTHING:
    - Convert vague statements to specific metrics
    - Add numbers, percentages, timeframes, volumes
    - Examples: "Processed samples" → "Processed 80-100 samples daily"

====================
OUTPUT FORMAT
====================

Create a complete, professionally formatted resume with these sections in order:

[CANDIDATE NAME - ALL CAPS]
[Contact line: Location | Email | Phone | LinkedIn]

PROFESSIONAL SUMMARY
[3-4 sentences, laser-focused on this job, incorporating exact job title and top requirements]

CORE COMPETENCIES
[List 8-12 most relevant skills from job description that candidate possesses, in bullet format]

WORK EXPERIENCE
[Most relevant role FIRST, then others in order of relevance]

For each role:
Company Name | Location | Dates
• [Bullet 1 - most relevant to job requirements]
• [Bullet 2 - second most relevant]
• [Bullet 3 - quantified achievement]
• [Additional bullets as needed]

EDUCATION
[Degree]
[Institution | Location | Dates]
[GPA if strong and recent]

CERTIFICATIONS & PROFESSIONAL DEVELOPMENT
[Relevant certifications, licenses, training]

[Optional: PROFESSIONAL ACHIEVEMENTS section if strong publications/awards]

====================
QUALITY CHECKLIST
====================

Before finalizing, verify:
✓ Most relevant past role is listed FIRST in experience
✓ Summary uses exact job title from posting
✓ Every major job requirement is addressed somewhere in resume
✓ No markdown formatting (no hashtags, asterisks for bold, etc.)
✓ No technical terms added that aren't in job posting or original resume
✓ Bullets prioritized by job relevance within each role
✓ Location match mentioned if applicable
✓ All dates and facts remain accurate
✓ Quantified achievements prominently featured
✓ ATS-friendly formatting throughout

====================
TONE & STYLE
====================

- Confident and accomplishment-focused
- Industry-specific terminology from job posting
- Active voice, strong action verbs
- Concise but comprehensive
- Professional without being stuffy
- Results-oriented language

Generate the complete resume now. Make it outstanding."""

    try:
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=6000,
            messages=[{"role": "user", "content": prompt}]
        )
        return message.content[0].text
    except Exception as e:
        st.error(f"Error generating resume: {str(e)}")
        return None

def create_docx_from_text(resume_text):
    """Create Word doc from resume text"""
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    lines = resume_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Remove any markdown formatting that might have slipped through
        line = line.replace('**', '').replace('##', '').replace('#', '')
        
        # Detect section headers (all caps or specific keywords)
        is_header = (
            line.isupper() and len(line) > 3 or
            line in ['PROFESSIONAL SUMMARY', 'CORE COMPETENCIES', 'TECHNICAL SKILLS', 
                    'WORK EXPERIENCE', 'EDUCATION', 'CERTIFICATIONS', 'PROFESSIONAL ACHIEVEMENTS',
                    'CERTIFICATIONS & PROFESSIONAL DEVELOPMENT', 'PROFESSIONAL DEVELOPMENT']
        )
        
        # Detect name (first non-empty line)
        is_name = len(doc.paragraphs) == 0 and not line.startswith('•') and not line.startswith('-')
        
        # Detect job titles (bold lines that aren't headers)
        is_job_title = (
            not is_header and 
            not line.startswith('•') and 
            not line.startswith('-') and
            not '@' in line and
            not '|' in line and
            len(line.split()) < 8 and
            any(word in line for word in ['Assistant', 'Specialist', 'Technician', 'Manager', 'Director', 'Coordinator', 'Analyst'])
        )
        
        para = doc.add_paragraph()
        run = para.add_run(line)
        
        if is_name:
            run.font.size = Pt(18)
            run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif is_header:
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(31, 56, 100)  # Dark blue for headers
            para.space_before = Pt(12)
            para.space_after = Pt(6)
        elif is_job_title:
            run.font.size = Pt(11)
            run.bold = True
            para.space_before = Pt(6)
        elif line.startswith('•') or line.startswith('-'):
            run.font.size = Pt(10)
            para.left_indent = Inches(0.25)
            para.paragraph_format.space_after = Pt(2)
        else:
            run.font.size = Pt(10)
    
    return doc

# Initialize session state
if 'api_key_validated' not in st.session_state:
    st.session_state.api_key_validated = False
if 'api_key' not in st.session_state:
    st.session_state.api_key = ""

# API Key Gate
if not st.session_state.api_key_validated:
    st.title("📄 AI Resume Tailor")
    st.markdown('<p class="subtitle">Generate ATS-optimized resumes in seconds</p>', unsafe_allow_html=True)
    
    api_key_input = st.text_input("Enter your Anthropic API Key", type="password", placeholder="sk-ant-api03-...")
    
    if st.button("🚀 Start Using App"):
        if api_key_input and api_key_input.startswith("sk-ant-"):
            try:
                client = anthropic.Anthropic(api_key=api_key_input)
                client.messages.create(model="claude-sonnet-4-20250514", max_tokens=10, messages=[{"role": "user", "content": "Hi"}])
                st.session_state.api_key = api_key_input
                st.session_state.api_key_validated = True
                st.rerun()
            except:
                st.error("❌ Invalid API key")
        else:
            st.error("❌ Enter valid API key")
    st.stop()

# Check if master resume exists
def get_master_resume_content():
    """Load and extract master resume content"""
    resume_data = None
    source = None
    
    # Try local file first
    if os.path.exists(MASTER_RESUME_PATH):
        try:
            with open(MASTER_RESUME_PATH, 'rb') as f:
                resume_data = f.read()
                source = "local file"
        except Exception as e:
            st.error(f"Error loading local resume: {str(e)}")
    
    # Try Streamlit secrets if local file not found
    if not resume_data and hasattr(st, 'secrets') and "master_resume_base64" in st.secrets:
        try:
            import base64
            resume_data = base64.b64decode(st.secrets["master_resume_base64"])
            source = "secrets"
        except Exception as e:
            st.error(f"Error loading resume from secrets: {str(e)}")
    
    # If no resume data found
    if not resume_data:
        st.error("❌ Master resume not found. Please upload master_resume.docx to your GitHub repository.")
        st.info("📁 The file should be in the same folder as the Python script.")
        return None
    
    # Extract text from resume data
    try:
        # Validate it's a proper Word document
        doc = Document(io.BytesIO(resume_data))
        
        # Extract all text
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        if not paragraphs:
            st.error("❌ Master resume appears to be empty. Please check the file.")
            return None
        
        content = '\n'.join(paragraphs)
        st.success(f"✅ Master resume loaded successfully from {source} ({len(paragraphs)} paragraphs, {len(content)} characters)")
        return content
        
    except Exception as e:
        st.error(f"❌ Error reading master resume: {str(e)}")
        st.info("💡 Make sure the file is a valid .docx (Microsoft Word) document, not .doc or PDF.")
        return None

# Main App
st.title("📄 Paula's Resume Generator")
st.markdown('<p class="subtitle">Generate ATS-optimized resumes in seconds</p>', unsafe_allow_html=True)

job_url = st.text_input("Job Posting URL *", placeholder="https://www.linkedin.com/jobs/view/...")
additional_context = st.text_area("Additional Context (Optional)", placeholder="E.g., 'Emphasize my laboratory skills' or 'Transitioning from QA to lab work'", height=100)

if st.button("🚀 Generate Tailored Resume"):
    if not job_url:
        st.error("Please enter job URL")
    else:
        # Create progress bar
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # Step 1: Analyze job posting
        status_text.markdown("**<span style='color: #1a202c;'>🔍 Step 1/3:</span>** <span style='color: #1a202c;'>Analyzing job posting...</span>", unsafe_allow_html=True)
        progress_bar.progress(10)
        
        job_content = scrape_job_posting(job_url)
        if not job_content:
            st.error("Failed to retrieve job posting. Please check the URL.")
            st.stop()
        
        progress_bar.progress(30)
        job_analysis = analyze_job_posting(job_content, st.session_state.api_key)
        
        if job_analysis:
            progress_bar.progress(40)
            status_text.markdown("**<span style='color: #1a202c;'>✅ Step 1/3:</span>** <span style='color: #1a202c;'>Job posting analyzed!</span>", unsafe_allow_html=True)
            
            # Step 2: Load master resume
            status_text.markdown("**<span style='color: #1a202c;'>📄 Step 2/3:</span>** <span style='color: #1a202c;'>Loading your master resume...</span>", unsafe_allow_html=True)
            progress_bar.progress(50)
            
            original_resume = get_master_resume_content()
            if not original_resume:
                st.error("❌ Master resume not found. Please add master_resume.docx to your repository.")
                st.stop()
            
            progress_bar.progress(60)
            status_text.markdown("**<span style='color: #1a202c;'>✅ Step 2/3:</span>** <span style='color: #1a202c;'>Master resume loaded!</span>", unsafe_allow_html=True)
            
            # Step 3: Generate tailored resume
            status_text.markdown("**<span style='color: #1a202c;'>✨ Step 3/3:</span>** <span style='color: #1a202c;'>Generating tailored resume (30-60 seconds)...</span>", unsafe_allow_html=True)
            progress_bar.progress(65)
            
            tailored_resume = generate_tailored_resume(original_resume, job_analysis, additional_context, st.session_state.api_key)
            
            if tailored_resume:
                progress_bar.progress(90)
                status_text.markdown("**<span style='color: #1a202c;'>✅ Step 3/3:</span>** <span style='color: #1a202c;'>Resume generated!</span>", unsafe_allow_html=True)
                
                # Create document
                doc = create_docx_from_text(tailored_resume)
                doc_bytes = io.BytesIO()
                doc.save(doc_bytes)
                doc_bytes.seek(0)
                
                progress_bar.progress(100)
                status_text.markdown("**<span style='color: #1a202c;'>🎉 Complete!</span>** <span style='color: #1a202c;'>Your resume is ready to download.</span>", unsafe_allow_html=True)
                
                st.success("✅ Resume generated successfully!")
                st.download_button(
                    "⬇️ Download Word Document",
                    doc_bytes.getvalue(),
                    f"Resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                st.balloons()
                
                # Clear progress indicators after success
                progress_bar.empty()
                status_text.empty()

# Features
st.markdown("""
<div class="feature-grid">
    <div class="feature-item">
        <div class="feature-icon">✓</div>
        <div class="feature-text">ATS Optimized</div>
    </div>
    <div class="feature-item">
        <div class="feature-icon">✓</div>
        <div class="feature-text">Keyword Matching</div>
    </div>
    <div class="feature-item">
        <div class="feature-icon">✓</div>
        <div class="feature-text">Smart Tailoring</div>
    </div>
    <div class="feature-item">
        <div class="feature-icon">✓</div>
        <div class="feature-text">Instant Results</div>
    </div>
</div>
""", unsafe_allow_html=True)
